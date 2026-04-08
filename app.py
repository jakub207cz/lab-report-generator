from __future__ import annotations

import io
import json
import mimetypes
import re
import unicodedata
import os
import math
import ast
import textwrap
import importlib
import time
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Set, Tuple, Union, Literal

import streamlit as st
import google.generativeai as genai
import pandas as pd
import pypdf
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pydantic import BaseModel, Field

# --- OPTIONAL DEPENDENCIES ---
try:
    from openpyxl import load_workbook
    from openpyxl.utils.cell import get_column_letter, range_boundaries
except Exception:
    load_workbook = None; get_column_letter = None; range_boundaries = None

try:
    from pdfminer.high_level import extract_text as pdfminer_extract_text
except Exception:
    pdfminer_extract_text = None

try:
    import pytesseract
except Exception:
    pytesseract = None

try:
    import pypdfium2 as pdfium
except Exception:
    pdfium = None

try:
    from pptx import Presentation
except Exception:
    Presentation = None

# --- MODELS / SCHEMAS (from pipeline/schemas.py) ---
class ImageAsset(BaseModel):
    image_id: str
    filename: str
    section: str

class LabReportData(BaseModel):
    teorie: str = Field(default="")
    postup: str = Field(default="")
    priklad_vypoctu: str = Field(default="")
    zaver: str = Field(default="")
    image_references: List[str] = Field(default_factory=list)

class QualityIssue(BaseModel):
    severity: Literal["WARN", "FAIL"]
    code: str
    message: str

class QualityGateResult(BaseModel):
    status: Literal["PASS", "WARN", "FAIL"]
    issues: List[QualityIssue] = Field(default_factory=list)

class DocumentChunk(BaseModel):
    text: str
    type: Literal["paragraph", "table", "figure", "heading"]
    source_file: str
    page: Optional[int] = None
    section_hint: Optional[str] = None
    confidence: float = 1.0

class TableData(BaseModel):
    headers: List[str]
    rows: List[List[Any]]
    units: Optional[List[str]] = None
    source_file: str
    page: Optional[int] = None
    sheet_name: Optional[str] = None
    section_hint: Optional[str] = None

class FigureData(BaseModel):
    figure_id: str
    source_file: str
    page: Optional[int] = None
    slide: Optional[int] = None
    section_hint: Optional[str] = None
    ocr_text: Optional[str] = None
    confidence: float = 1.0

class NormalizedIngestionResult(BaseModel):
    chunks: List[DocumentChunk] = Field(default_factory=list)
    tables: List[TableData] = Field(default_factory=list)
    figures: List[FigureData] = Field(default_factory=list)
    metadata: Dict[str, Any] = Field(default_factory=dict)

# --- INGESTION PIPELINE (from pipeline/ingestion.py) ---
UNIT_PATTERN = re.compile(r"(?:\(([^)]+)\)|\[([^\]]+)\])")
ASSIGNMENT_SECTION_ALIASES: List[tuple[str, str, List[str]]] = [
    ("assignment", "Zadání", ["zadani", "zadání"]),
    ("theory", "Teoretický úvod", ["teoreticky uvod", "teorie"]),
    ("schema", "Schéma zapojení", ["schema zapojeni", "schéma zapojení"]),
    ("procedure", "Postup měření", ["postup mereni", "postup měření"]),
    ("tables_example", "Příklad tabulek", ["priklad tabulek", "tabulky"]),
    ("calculation_example", "Příklad výpočtů", ["priklad vypoctu", "vypocty"]),
    ("conclusion", "Závěr", ["zaver", "závěr"]),
]

@dataclass
class BinarySource:
    filename: str
    data: bytes
    section_hint: Optional[str] = None
    mime_type: Optional[str] = None

class IngestionPipeline:
    def __init__(self, enable_ocr: bool = True) -> None:
        self.enable_ocr = enable_ocr
        self._figure_counter = 1

    def ingest_sources(self, sources: Iterable[BinarySource]) -> NormalizedIngestionResult:
        res = NormalizedIngestionResult(); meta = {}
        for s in sources:
            p = self._resolve_parser(Path(s.filename).suffix.lower())
            if not p: continue
            out = p(s)
            res.chunks.extend(out.chunks); res.tables.extend(out.tables); res.figures.extend(out.figures)
            for k, v in out.metadata.items():
                if k not in meta: meta[k] = v
                elif isinstance(meta[k], dict) and isinstance(v, dict): meta[k].update(v)
                elif isinstance(meta[k], list) and isinstance(v, list): meta[k].extend(v)
        res.metadata = {"source_count": len(list(sources)), "ocr_enabled": self.enable_ocr}
        res.metadata.update(meta); return res

    def ingest_streamlit_files(self, files: Iterable[Any], hint: Optional[str] = None) -> NormalizedIngestionResult:
        srcs = [BinarySource(f.name, f.getvalue(), hint, getattr(f, "type", None)) for f in (files or [])]
        return self.ingest_sources(srcs)

    def _resolve_parser(self, ext: str):
        if ext == ".docx": return self.docx_parser
        if ext in {".xlsx", ".xls"}: return self.xlsx_parser
        if ext in {".txt", ".csv"}: return self.text_parser
        if ext == ".pdf": return self.pdf_parser
        if ext == ".pptx": return self.pptx_parser
        if ext in {".png", ".jpg", ".jpeg"}: return self.image_handler
        return None

    def _next_figure_id(self) -> str:
        fid = f"FIG-{self._figure_counter:03d}"; self._figure_counter += 1; return fid

    def _units_from_headers(self, headers: List[str]) -> List[str]:
        return [(UNIT_PATTERN.search(h).group(1) or UNIT_PATTERN.search(h).group(2)).strip() if UNIT_PATTERN.search(h or "") else "" for h in headers]

    def docx_parser(self, s: BinarySource) -> NormalizedIngestionResult:
        doc, out = Document(io.BytesIO(s.data)), NormalizedIngestionResult()
        for p in doc.paragraphs:
            if p.text.strip(): out.chunks.append(DocumentChunk(text=p.text.strip(), type="heading" if "heading" in (p.style.name or "").lower() else "paragraph", source_file=s.filename, section_hint=s.section_hint))
        for t in doc.tables:
            rows = [[(c.text or "").strip() for c in r.cells] for r in t.rows]
            if rows: out.tables.append(TableData(headers=rows[0], rows=rows[1:], source_file=s.filename, section_hint=s.section_hint))
        return out

    def text_parser(self, s: BinarySource) -> NormalizedIngestionResult:
        out = NormalizedIngestionResult(); t = ""
        for e in ("utf-8", "cp1250"):
            try: t = s.data.decode(e).strip(); break
            except: t = ""
        if t:
            out.chunks.append(DocumentChunk(text=t, type="paragraph", source_file=s.filename, section_hint=s.section_hint))
            if s.filename.lower().endswith(".csv"):
                try: df = pd.read_csv(io.BytesIO(s.data)); out.tables.append(TableData(headers=list(df.columns), rows=df.fillna("").values.tolist(), source_file=s.filename, section_hint=s.section_hint))
                except: pass
        return out

    def xlsx_parser(self, s: BinarySource) -> NormalizedIngestionResult:
        out = NormalizedIngestionResult()
        if load_workbook:
            try:
                wb = load_workbook(io.BytesIO(s.data), data_only=True)
                for ws in wb.worksheets:
                    data = [[c.value for c in r] for r in ws.iter_rows()]
                    if data and any(any(c is not None for c in r) for r in data):
                        out.tables.append(TableData(headers=[str(v or "") for v in data[0]], rows=[[str(v or "") for v in r] for r in data[1:]], source_file=s.filename, sheet_name=ws.title, section_hint=s.section_hint))
            except: pass
        return out

    def pdf_parser(self, s: BinarySource) -> NormalizedIngestionResult:
        out = NormalizedIngestionResult()
        try:
            r = pypdf.PdfReader(io.BytesIO(s.data))
            for i, p in enumerate(r.pages, 1):
                t = (p.extract_text() or "").strip()
                if t: out.chunks.append(DocumentChunk(text=t, type="paragraph", source_file=s.filename, page=i, section_hint=s.section_hint))
        except: pass
        return out

    def image_handler(self, s: BinarySource) -> NormalizedIngestionResult:
        out = NormalizedIngestionResult(); fid = self._next_figure_id()
        out.figures.append(FigureData(figure_id=fid, source_file=s.filename, section_hint=s.section_hint))
        return out

# --- GENERATION (from pipeline/generation.py) ---
def _strip_markdown_fences(raw_text: str) -> str:
    text = raw_text.strip()
    if "```json" in text: text = text.split("```json", 1)[1].split("```", 1)[0]
    elif "```" in text: text = text.split("```", 1)[1].split("```", 1)[0]
    return text.strip()

def _extract_json_object(raw_text: str) -> str:
    text = _strip_markdown_fences(raw_text); start, end = text.find("{"), text.rfind("}")
    return text[start : end + 1] if start >= 0 and end > start else text

class _TheoryProcedureData(BaseModel):
    teorie: str = Field(default=""); postup: str = Field(default=""); image_references: List[str] = Field(default_factory=list)

class _ConclusionData(BaseModel):
    zaver: str = Field(default=""); image_references: List[str] = Field(default_factory=list)

class _CalculationItem(BaseModel):
    title: str = Field(default=""); general_formula_latex: str = Field(default=""); substitution_formula_latex: str = Field(default=""); compute_expression: str = Field(default=""); result_symbol_latex: str = Field(default="x"); result_unit_latex: str = Field(default=""); variables: Dict[str, float] = Field(default_factory=dict)

class _CalculationData(BaseModel):
    items: List[_CalculationItem] = Field(default_factory=list)

def _generate_structured_part(model, prompt, schema_model) -> BaseModel:
    for attempt in range(3):
        try:
            res = model.generate_content([prompt])
            raw = _extract_json_object(res.text)
            return schema_model.model_validate(json.loads(raw))
        except Exception as e:
            if "429" in str(e) or "ResourceExhausted" in str(e):
                if attempt < 2:
                    st.warning(f"⚠️ Dosáhli jste limitu požadavků. Čekám 10 sekund... (pokus {attempt+1}/3)")
                    time.sleep(10)
                    continue
                else:
                    st.error("❌ Google API je přetížené nebo jste vyčerpali limit. Zkuste to za minutu znovu.")
                    raise e
            try:
                repair = model.generate_content([f"Fix JSON for schema {json.dumps(schema_model.model_json_schema())}:\n{res.text}"])
                return schema_model.model_validate(json.loads(_extract_json_object(repair.text)))
            except:
                if attempt == 2: raise e
                time.sleep(2)
    return None

def generate_lab_report_advanced(api_key, model_name, topic, inputs_map, is_handwritten=False) -> LabReportData:
    genai.configure(api_key=api_key)
    try:
        model = genai.GenerativeModel(model_name)
        # Test call to verify model availability
        model.generate_content("test", generation_config={"max_output_tokens": 1})
    except:
        st.info("⚠️ Vybraný model není dostupný, používám stabilní verzi gemini-1.5-flash.")
        model = genai.GenerativeModel("gemini-1.5-flash")
    
    t_len = "ZKRÁCENÝ ROZSAH: Max půl strany A4!" if is_handwritten else "Max 1.5 strany A4."
    c_len = "ZKRÁCENÝ ROZSAH: Max půl strany A4!" if is_handwritten else "Fakta a detailní analýza"
    
    t_prompt = f"Téma: {topic}\n\nTEORIE ({t_len})\nPodklady: {inputs_map.get('theory_text', '')}\nPOSTUP\nPodklady: {inputs_map.get('procedure_text', '')}\nSchema: {json.dumps(_TheoryProcedureData.model_json_schema())}"
    c_prompt = f"Téma: {topic}\n\nZÁVĚR ({c_len})\nPodklady: {inputs_map.get('conclusion_text', '')}\nData: {inputs_map.get('data_text', '')}\nSchema: {json.dumps(_ConclusionData.model_json_schema())}"
    cal_prompt = f"Téma: {topic}\n\nVytvoř 3-4 výpočty. Schema: {json.dumps(_CalculationData.model_json_schema())}"

    t_data = _generate_structured_part(model, t_prompt, _TheoryProcedureData)
    time.sleep(2) # Delay to prevent rate limiting
    c_data = _generate_structured_part(model, c_prompt, _ConclusionData)
    time.sleep(2)
    try: 
        cal_data = _generate_structured_part(model, cal_prompt, _CalculationData)
    except: 
        cal_data = _CalculationData()

    return LabReportData(
        teorie=t_data.teorie if t_data else "Chyba generování.", 
        postup=t_data.postup if t_data else "Chyba generování.", 
        zaver=c_data.zaver if c_data else "Chyba generování.", 
        priklad_vypoctu=json.dumps(cal_data.model_dump(), ensure_ascii=False) if cal_data else "{}", 
        image_references=list(set((t_data.image_references if t_data else []) + (c_data.image_references if c_data else [])))
    )

# --- DOCX WRITER (from render/docx_writer.py) ---
def _render_eq(latex):
    try:
        import matplotlib.pyplot as plt
        fig = plt.figure(figsize=(8, 1.25)); ax = fig.add_axes([0,0,1,1]); ax.axis("off")
        ax.text(0.02, 0.5, f"${latex}$", fontsize=20, va="center", ha="left")
        bio = io.BytesIO(); fig.savefig(bio, format="png", transparent=True, bbox_inches="tight"); plt.close(fig); bio.seek(0); return bio
    except: return None

def fill_template_docx_advanced(template_path, inputs_map, ai_content: LabReportData, topic=""):
    doc = Document(template_path) if os.path.exists(template_path) else Document()
    topic = (topic or str(inputs_map.get("topic", ""))).strip()
    if doc.tables:
        t = doc.tables[0]
        # Vyhledání a doplnění názvu úlohy
        for r in t.rows:
            if "název úlohy" in (r.cells[0].text or "").lower():
                r.cells[1].text = topic
                for p in r.cells[1].paragraphs: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    sec_map = {
        "Teoretický úvod": {"text": ai_content.teorie, "imgs": inputs_map.get("waveforms_images", [])},
        "Schéma zapojení": {"text": "", "imgs": inputs_map.get("schema_images", [])},
        "Postup měření": {"text": ai_content.postup, "imgs": []},
        "Naměřené a vypočítané hodnoty": {"text": "", "imgs": inputs_map.get("data_images", []), "tables": inputs_map.get("data_tables", [])},
        "Příklad výpočtu": {"text": ai_content.priklad_vypoctu, "imgs": [], "is_calc": True},
        "Závěr": {"text": ai_content.zaver, "imgs": []}
    }

    def add_c(para, data):
        parent, idx = para._element.getparent(), para._element.getparent().index(para._element)
        if data.get("text") and not data.get("is_calc"):
            p = doc.add_paragraph(data["text"]); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            parent.insert(idx + 1, p._element); idx += 1
        if data.get("is_calc"):
            try:
                calc = json.loads(data["text"])
                for item in calc.get("items", []):
                    p = doc.add_paragraph(item["title"]); p.runs[0].bold = True; parent.insert(idx + 1, p._element); idx += 1
                    for lat in [item["general_formula_latex"], item["substitution_formula_latex"]]:
                        img = _render_eq(lat)
                        p = doc.add_paragraph()
                        if img: p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.add_run().add_picture(img, width=Pt(360))
                        else: p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.add_run(lat)
                        parent.insert(idx + 1, p._element); idx += 1
            except: pass
        for img in data.get("imgs", []):
            try:
                b = io.BytesIO(); img.save(b, "PNG"); b.seek(0)
                p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.add_run().add_picture(b, width=Pt(400)); parent.insert(idx + 1, p._element); idx += 1
            except: pass
        for tab in data.get("tables", []):
            try:
                t = doc.add_table(rows=len(tab["rows"])+1, cols=len(tab["headers"])); t.style = "Table Grid"
                for i, h in enumerate(tab["headers"]): t.rows[0].cells[i].text = str(h)
                for ri, r in enumerate(tab["rows"]):
                    for ci, v in enumerate(r): t.rows[ri+1].cells[ci].text = str(v)
                parent.insert(idx + 1, t._element); idx += 1
            except: pass

    for p in list(doc.paragraphs):
        txt = p.text.strip()
        if txt in sec_map: add_c(p, sec_map[txt])
    
    bio = io.BytesIO(); doc.save(bio); return bio

# --- STREAMLIT UI ---
st.set_page_config(page_title="AI Lab Report Generator", layout="wide", page_icon="⚡")
st.markdown("""
<style>
    .main-title { text-align: center; font-size: clamp(2.5em, 6vw, 4em); font-weight: 800; background: linear-gradient(45deg, #00f2fe, #4facfe, #00f2fe); -webkit-background-clip: text; -webkit-text-fill-color: transparent; }
    .sub-title { text-align: center; color: #a0aec0; font-size: 1.2em; margin-bottom: 2em; }
    .stButton>button { background: linear-gradient(90deg, #00f2fe, #4facfe); color: #0e1117 !important; border-radius: 12px; font-weight: 800; padding: 14px; width: 100%; }
</style>
""", unsafe_allow_html=True)

st.markdown("<h1 class='main-title'>⚡ AI Generátor Protokolů</h1>", unsafe_allow_html=True)
st.markdown("<p class='sub-title'>Tvoje záchrana pro laborky na SPŠE. Rychle a bez stresu.</p>", unsafe_allow_html=True)

with st.expander("🔑 Nastavení & API", expanded=True):
    col1, col2 = st.columns(2)
    with col1: api_key = st.text_input("Google Gemini API Key", type="password")
    with col2:
        model_options = {
            "Gemini 2.5 Flash (Rychlý)": "gemini-2.5-flash",
            "Gemini 3.1 Flash (Nejnovější)": "gemini-3.1-flash"
        }
        selected_model_label = st.radio("Vyberte model AI:", options=list(model_options.keys()), index=0)
        model_choice = model_options[selected_model_label]

with st.form("lab_report_form"):
    topic = st.text_input("Téma měření", placeholder="Např. Oživování a měření na stabilizovaném zdroji...")
    is_handwritten = st.toggle("📝 Píšu tenhle elaborát ručně! (Zkrátit texty)")
    
    st.markdown("### 📂 Podklady pro AI")
    asgn_f = st.file_uploader("Zadání úlohy", accept_multiple_files=True, key="asgn")
    data_f = st.file_uploader("Tabulky naměřených hodnot", accept_multiple_files=True, key="data")
    theo_f = st.file_uploader("Podklady k teorii", accept_multiple_files=True, key="theo")
    wave_f = st.file_uploader("Grafické průběhy", accept_multiple_files=True, key="wave")
    proc_f = st.file_uploader("Pracovní postup", accept_multiple_files=True, key="proc")
    concl_f = st.file_uploader("Osnova pro závěr", accept_multiple_files=True, key="concl")
    schm_f = st.file_uploader("Schéma zapojení (obrázky)", accept_multiple_files=True, key="schm")
    
    submitted = st.form_submit_button("🚀 Vygenerovat nadupaný protokol")

if submitted:
    if not api_key or not topic: st.error("❌ Doplň API klíč a téma!")
    else:
        with st.spinner("Zpracovávám a generuji..."):
            pipeline = IngestionPipeline()
            asgn_r = pipeline.ingest_streamlit_files(asgn_f, "assignment")
            data_r = pipeline.ingest_streamlit_files(data_f, "data")
            theo_r = pipeline.ingest_streamlit_files(theo_f, "theory")
            wave_r = pipeline.ingest_streamlit_files(wave_f, "waveforms")
            proc_r = pipeline.ingest_streamlit_files(proc_f, "procedure")
            concl_r = pipeline.ingest_streamlit_files(concl_f, "conclusion")
            
            def get_txt(r): return "\n".join([c.text for c in r.chunks])
            def get_imgs(files):
                imgs = []
                for f in (files or []):
                    if f.type.startswith('image'): imgs.append(Image.open(f))
                return imgs

            inputs = {
                'topic': topic,
                'assignment_text': get_txt(asgn_r),
                'data_text': get_txt(data_r),
                'theory_text': get_txt(theo_r),
                'procedure_text': get_txt(proc_r),
                'conclusion_text': get_txt(concl_r),
                'waveforms_images': get_imgs(wave_f),
                'schema_images': get_imgs(schm_f),
                'data_images': get_imgs(data_f),
                'data_tables': [{"headers": t.headers, "rows": t.rows} for t in data_r.tables]
            }
            
            ai_data = generate_lab_report_advanced(api_key, model_choice, topic, inputs, is_handwritten)
            
            st.balloons(); st.success("🎉 Hotovo!")
            with st.expander("Náhled", expanded=True):
                st.markdown(f"**Teorie:**\n{ai_data.teorie[:500]}...")
            
            doc_file = fill_template_docx_advanced("Graficka_Osnova.docx", inputs, ai_data, topic)
            st.download_button("📥 Stáhnout protokol (.docx)", doc_file.getvalue(), "laboratorni_protokol.docx")
