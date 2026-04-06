from __future__ import annotations

import io
import importlib
import ast
import json
import math
import os
import re
import textwrap
from typing import Any, Dict, List

import streamlit as st
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from pipeline.schemas import ImageAsset, LabReportData


MAX_TABLE_COLUMNS_PER_BLOCK = 8

_SUPPRESSED_TEXT_LINE_PATTERNS = [
    re.compile(r"^\s*Detekované\s+grafy\s+z\s+XLSX\s*:\s*$", flags=re.IGNORECASE),
    re.compile(r"^\s*-\s*FIG-\d+\s*:\s*.*$", flags=re.IGNORECASE),
    re.compile(r"^\s*Tabulky\s+z\s+XLSX\s+byly\s+vloženy\s+níže\s+jako\s+skutečné\s+tabulky\.?\s*$", flags=re.IGNORECASE),
]

_PLOTTER = None


def _find_row_by_prefix(table, prefix: str, col_idx: int = 0) -> int:
    prefix_norm = (prefix or "").strip().lower()
    for i, row in enumerate(table.rows):
        if col_idx >= len(row.cells):
            continue
        cell_text = (row.cells[col_idx].text or "").strip().lower()
        if cell_text.startswith(prefix_norm):
            return i
    return -1


def _estimate_sheet_count(inputs_map: dict, ai_content: LabReportData) -> int:
    text_blocks = [
        ai_content.teorie,
        ai_content.postup,
        ai_content.priklad_vypoctu,
        ai_content.zaver,
        str(inputs_map.get("instruments_text", "")),
        str(inputs_map.get("assignment_text", "")),
    ]
    total_chars = sum(len((block or "").strip()) for block in text_blocks)

    images_count = (
        len(inputs_map.get("schema_images", []))
        + len(inputs_map.get("waveforms_images", []))
        + len(inputs_map.get("data_images", []))
    )

    # 1 stránka = desky, další orientačně dle textu + obrázků
    text_pages = max(1, (total_chars // 2800) + (1 if total_chars % 2800 else 0)) if total_chars else 1
    image_pages = (images_count // 2) + (1 if images_count % 2 else 0)
    return max(1, 1 + text_pages + image_pages)


def _prepare_assignment_lines_for_cover(assignment_text: str) -> List[str]:
    raw = (assignment_text or "").replace("\r\n", "\n").replace("\r", "\n")
    if not raw.strip():
        return []

    cleaned_lines: List[str] = []
    for line in raw.split("\n"):
        stripped = line.strip()
        if not stripped:
            cleaned_lines.append("")
            continue
        if stripped.startswith("---") and stripped.endswith("---"):
            continue
        if re.fullmatch(r"[A-Z0-9_\- ]+\.(DOC|DOCX|PDF)", stripped, flags=re.IGNORECASE):
            continue
        if re.fullmatch(r"\d{1,3}", stripped):
            continue
        cleaned_lines.append(stripped)

    cleaned_text = "\n".join(cleaned_lines).strip()
    if not cleaned_text:
        return []

    zadani_match = re.search(r"\bzadání\s*:\s*", cleaned_text, flags=re.IGNORECASE)
    if zadani_match:
        cleaned_text = cleaned_text[zadani_match.end():].strip()

    stop_match = re.search(
        r"\b(teoretický\s+úvod|schéma\s+zapojení|schema\s+zapojeni|postup\s+měření|postup\s+mereni)\b",
        cleaned_text,
        flags=re.IGNORECASE,
    )
    if stop_match:
        cleaned_text = cleaned_text[:stop_match.start()].strip()

    compact_text = re.sub(r"\s*\n\s*", " ", cleaned_text)
    compact_text = re.sub(r"\s{2,}", " ", compact_text).strip()
    if not compact_text:
        return []

    numbered_items = [
        re.sub(r"\s{2,}", " ", item).strip()
        for item in re.findall(r"\d+\.\s.*?(?=(?:\s\d+\.\s)|$)", compact_text)
    ]
    if numbered_items:
        return numbered_items

    return [compact_text]


def _set_cover_page_fields(doc: Document, topic: str, username: str, assignment_text: str, sheet_count: int) -> None:
    if not doc.tables:
        return

    cover_table = doc.tables[0]

    task_row = _find_row_by_prefix(cover_table, "Název úlohy", col_idx=0)
    if task_row >= 0:
        target_row = min(task_row + 1, len(cover_table.rows) - 1)
        if len(cover_table.rows[target_row].cells) > 1:
            title_cell = cover_table.rows[target_row].cells[1]
            title_cell.text = (topic or "").strip()
            if title_cell.paragraphs:
                title_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    author_row = _find_row_by_prefix(cover_table, "Vypracoval", col_idx=0)
    if author_row >= 0:
        target_row = min(author_row + 1, len(cover_table.rows) - 1)
        if len(cover_table.rows[target_row].cells) > 1:
            author_cell = cover_table.rows[target_row].cells[1]
            author_cell.text = (username or "").strip()
            if author_cell.paragraphs:
                author_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for i, row in enumerate(cover_table.rows):
        for cell_idx, cell in enumerate(row.cells):
            if "počet listů" in (cell.text or "").strip().lower():
                target_row = min(i + 1, len(cover_table.rows) - 1)
                if cell_idx < len(cover_table.rows[target_row].cells):
                    sheets_cell = cover_table.rows[target_row].cells[cell_idx]
                    sheets_cell.text = str(sheet_count)
                    if sheets_cell.paragraphs:
                        sheets_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                break

    zadani_anchor = -1
    for i, row in enumerate(cover_table.rows):
        if any("zadání" in (cell.text or "").strip().lower() for cell in row.cells):
            zadani_anchor = i
            break

    if zadani_anchor >= 0:
        start_row = min(zadani_anchor + 1, len(cover_table.rows) - 1)
        end_row = len(cover_table.rows)

        prepared_lines = _prepare_assignment_lines_for_cover(assignment_text)
        if not prepared_lines:
            prepared_lines = []
            for paragraph in (assignment_text or "").splitlines():
                p = paragraph.strip()
                if not p:
                    continue
                prepared_lines.extend(textwrap.wrap(p, width=75) or [p])

        max_lines = max(0, end_row - start_row)
        fitted_lines = prepared_lines[:max_lines]
        if prepared_lines and len(prepared_lines) > max_lines and max_lines > 0:
            fitted_lines[-1] = fitted_lines[-1][:72] + "..."

        for row_idx in range(start_row, end_row):
            if row_idx - start_row < len(fitted_lines):
                cover_table.rows[row_idx].cells[0].text = fitted_lines[row_idx - start_row]
            else:
                cover_table.rows[row_idx].cells[0].text = ""


def _set_document_header_student(doc: Document, username: str) -> None:
    username_clean = (username or "").strip()

    def _append_page_number_field(paragraph) -> None:
        run = paragraph.add_run()
        fld_char_begin = OxmlElement("w:fldChar")
        fld_char_begin.set(qn("w:fldCharType"), "begin")
        run._r.append(fld_char_begin)

        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = "PAGE"
        run._r.append(instr_text)

        fld_char_end = OxmlElement("w:fldChar")
        fld_char_end.set(qn("w:fldCharType"), "end")
        run._r.append(fld_char_end)

    for section in doc.sections:
        header = section.header
        right_tab_position = section.page_width - section.left_margin - section.right_margin

        if header.paragraphs:
            paragraph = header.paragraphs[0]
        else:
            paragraph = header.add_paragraph()

        paragraph.text = ""
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph_format = paragraph.paragraph_format
        paragraph_format.left_indent = 0
        paragraph_format.right_indent = 0
        paragraph_format.first_line_indent = 0

        tab_stops = paragraph_format.tab_stops
        if hasattr(tab_stops, "clear_all"):
            tab_stops.clear_all()
        tab_stops.add_tab_stop(right_tab_position, WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES)

        if username_clean:
            paragraph.add_run(username_clean)
            paragraph.add_run("\t")
        _append_page_number_field(paragraph)


def _sanitize_section_text(text: str) -> str:
    if not text:
        return ""

    filtered_lines: List[str] = []
    for line in text.splitlines():
        stripped = line.strip()
        if any(pattern.match(stripped) for pattern in _SUPPRESSED_TEXT_LINE_PATTERNS):
            continue
        filtered_lines.append(line.rstrip())

    sanitized = "\n".join(filtered_lines)
    sanitized = re.sub(r"\n{3,}", "\n\n", sanitized)
    return sanitized.strip()


def _safe_float(value: Any) -> float | None:
    if value is None:
        return None
    if isinstance(value, (int, float)):
        return float(value)
    raw = str(value).strip().replace(" ", "")
    if not raw:
        return None
    if re.fullmatch(r"[+-]?\d+(?:[\.,]\d+)?(?:[eE][+-]?\d+)?", raw):
        try:
            return float(raw.replace(",", "."))
        except Exception:
            return None
    return None


def _fmt_num(value: float, max_decimals: int = 6) -> str:
    rendered = f"{value:.{max_decimals}f}".rstrip("0").rstrip(".")
    if rendered in {"", "-0"}:
        return "0"
    return rendered


def _compact_numeric_literals(text: str, max_decimals: int = 3) -> str:
    if not text:
        return ""

    number_pattern = re.compile(r"(?<![A-Za-z_])([+-]?\d+(?:[\.,]\d+)?(?:[eE][+-]?\d+)?)(?![A-Za-z_])")

    def _replace(match: re.Match[str]) -> str:
        raw = match.group(1)
        normalized = raw.replace(",", ".")
        try:
            num = float(normalized)
        except Exception:
            return raw
        return _fmt_num(num, max_decimals=max_decimals)

    return number_pattern.sub(_replace, text)


def _normalize_result_unit_latex(unit: str) -> str:
    normalized = (unit or "").strip()
    if not normalized:
        return ""

    normalized = re.sub(r"\\text\{([^}]*)\}", r"\1", normalized)
    normalized = re.sub(r"\\mathrm\{([^}]*)\}", r"\1", normalized)
    normalized = normalized.replace("\\", "")
    return normalized.strip()


def _normalize_latex_line(line: str) -> str:
    normalized = (line or "").strip()
    if not normalized:
        return ""

    # LLM často vrací escapovaný LaTeX z JSON (\\frac); pro renderer potřebujeme \frac.
    normalized = normalized.replace("\\\\", "\\")

    # Typicky nevalidní / nekompaktní konstrukce z LLM.
    normalized = re.sub(r"\\mathrm\{\\text\{([^}]*)\}\}", r"\\mathrm{\1}", normalized)
    normalized = re.sub(r"\\text\{([^}]*)\}", r"\1", normalized)
    normalized = _compact_numeric_literals(normalized, max_decimals=3)
    return normalized


_ALLOWED_MATH_FUNCS = {
    "sqrt": math.sqrt,
    "sin": math.sin,
    "cos": math.cos,
    "tan": math.tan,
    "log": math.log,
    "exp": math.exp,
    "abs": abs,
}

_ALLOWED_MATH_CONSTS = {
    "pi": math.pi,
    "e": math.e,
}


def _safe_eval_expression(expression: str, variables: Dict[str, float]) -> float | None:
    expr = (expression or "").strip().replace("^", "**")
    if not expr:
        return None

    try:
        parsed = ast.parse(expr, mode="eval")
    except Exception:
        return None

    allowed_nodes = (
        ast.Expression,
        ast.BinOp,
        ast.UnaryOp,
        ast.Constant,
        ast.Name,
        ast.Load,
        ast.Call,
        ast.Add,
        ast.Sub,
        ast.Mult,
        ast.Div,
        ast.Pow,
        ast.Mod,
        ast.USub,
        ast.UAdd,
    )

    for node in ast.walk(parsed):
        if not isinstance(node, allowed_nodes):
            return None
        if isinstance(node, ast.Call):
            if not isinstance(node.func, ast.Name):
                return None
            if node.func.id not in _ALLOWED_MATH_FUNCS:
                return None
        if isinstance(node, ast.Name):
            if node.id not in variables and node.id not in _ALLOWED_MATH_FUNCS and node.id not in _ALLOWED_MATH_CONSTS:
                return None

    safe_names: Dict[str, Any] = {}
    safe_names.update(_ALLOWED_MATH_FUNCS)
    safe_names.update(_ALLOWED_MATH_CONSTS)
    safe_names.update(variables)

    try:
        result = eval(compile(parsed, "<calc_expr>", "eval"), {"__builtins__": {}}, safe_names)
        return float(result)
    except Exception:
        return None


def _render_equation_image(latex: str) -> io.BytesIO | None:
    global _PLOTTER

    if _PLOTTER is None:
        try:
            matplotlib = importlib.import_module("matplotlib")
            matplotlib.use("Agg")
            _PLOTTER = importlib.import_module("matplotlib.pyplot")
        except Exception:
            _PLOTTER = False

    if _PLOTTER is False:
        return None

    plt = _PLOTTER
    fig = plt.figure(figsize=(8, 1.25), dpi=220)
    try:
        fig.patch.set_alpha(0)
        ax = fig.add_axes([0, 0, 1, 1])
        ax.axis("off")
        ax.text(0.02, 0.5, f"${latex}$", fontsize=20, va="center", ha="left")

        bio = io.BytesIO()
        fig.savefig(bio, format="png", transparent=True, bbox_inches="tight", pad_inches=0.05)
        bio.seek(0)
        return bio
    except Exception:
        return None
    finally:
        plt.close(fig)


def _build_calculation_equations(calculation_payload_text: str) -> List[Dict[str, str]]:
    if not calculation_payload_text:
        return []

    try:
        payload = json.loads(calculation_payload_text)
    except Exception:
        return []

    items = payload.get("items") if isinstance(payload, dict) else None
    if not isinstance(items, list):
        return []

    equations: List[Dict[str, str]] = []
    for item in items:
        if not isinstance(item, dict):
            continue

        title = str(item.get("title") or "").strip()
        general = _normalize_latex_line(str(item.get("general_formula_latex") or ""))
        substitution = _normalize_latex_line(str(item.get("substitution_formula_latex") or ""))
        expression = str(item.get("compute_expression") or "").strip()
        result_symbol = str(item.get("result_symbol_latex") or "x").strip()
        result_unit = _normalize_result_unit_latex(str(item.get("result_unit_latex") or ""))
        raw_variables = item.get("variables") if isinstance(item.get("variables"), dict) else {}

        if not (title and general and substitution and expression):
            continue

        variables: Dict[str, float] = {}
        for key, value in raw_variables.items():
            key_norm = str(key or "").strip()
            if not key_norm:
                continue
            as_float = _safe_float(value)
            if as_float is None:
                continue
            variables[key_norm] = as_float

        result_value = _safe_eval_expression(expression, variables)
        if result_value is None:
            continue

        result_line = f"{result_symbol} = {_fmt_num(result_value, max_decimals=3)}"
        if result_unit:
            result_line += rf"\,\mathrm{{{result_unit}}}"

        equations.append(
            {
                "title": title,
                "general": general,
                "substitution": substitution,
                "result": result_line,
            }
        )

    return equations


def fill_template_docx(
    template_path: str,
    inputs_map: dict,
    ai_content: LabReportData,
    image_registry: Dict[str, ImageAsset],
    username: str = "",
    topic: str = "",
) -> io.BytesIO:
    if not os.path.exists(template_path):
        st.error(f"Šablona {template_path} nenalezena!")
        doc = Document()
    else:
        doc = Document(template_path)

    resolved_username = (username or str(inputs_map.get("username", ""))).strip()
    resolved_topic = (topic or str(inputs_map.get("topic", ""))).strip()
    resolved_assignment = str(inputs_map.get("assignment_text", "")).strip()
    report_scope = str(inputs_map.get("report_scope", "full") or "full").strip().lower()

    sheet_count = _estimate_sheet_count(inputs_map, ai_content)
    _set_cover_page_fields(doc, resolved_topic, resolved_username, resolved_assignment, sheet_count)
    _set_document_header_student(doc, resolved_username)

    is_preparation = report_scope == "preparation"
    is_ending = report_scope == "ending"

    sections_map = {
        "Teoretický úvod": {"text": ai_content.teorie if not is_ending else "", "images": [], "image_ids": []},
        "Schéma zapojení": {
            "text": "",
            "images": inputs_map.get("schema_images", []) if not is_ending else [],
            "image_ids": inputs_map.get("schema_image_ids", []) if not is_ending else [],
        },
        "Postup měření": {"text": ai_content.postup if not is_ending else "", "images": [], "image_ids": []},
        "Naměřené a vypočítané hodnoty": {
            "text": str(inputs_map.get("data_text", "")) if not is_preparation else "",
            "images": inputs_map.get("data_images", []) if not is_preparation else [],
            "image_ids": inputs_map.get("data_image_ids", []) if not is_preparation else [],
            "tables": inputs_map.get("data_tables", []) if not is_preparation else [],
        },
        "Příklad výpočtu": {
            "text": ai_content.priklad_vypoctu if report_scope == "full" else "",
            "images": [],
            "image_ids": [],
            "equation_mode": report_scope == "full",
        },
        "Příklad výpočtů": {
            "text": ai_content.priklad_vypoctu if report_scope == "full" else "",
            "images": [],
            "image_ids": [],
            "equation_mode": report_scope == "full",
        },
        "Soupis použitých přístrojů": {
            "text": str(inputs_map.get("instruments_text", "")) if not is_preparation else "",
            "images": [],
            "image_ids": [],
        },
        "Grafy": {
            "text": str(inputs_map.get("waveforms_text", "")) if not is_preparation else "",
            "images": inputs_map.get("waveforms_images", []) if not is_preparation else [],
            "image_ids": inputs_map.get("waveforms_image_ids", []) if not is_preparation else [],
            "one_image_per_page": True,
            "show_image_caption": False,
        },
        "Závěr": {"text": ai_content.zaver if not is_preparation else "", "images": [], "image_ids": []},
    }

    def _format_cell_value(value: Any, round_numeric: bool) -> str:
        text = str(value) if value is not None else ""
        if not round_numeric:
            return text

        normalized = text.strip().replace(" ", "")
        if not normalized:
            return ""

        # Zaokrouhlíme jen čistě numerické hodnoty.
        if re.fullmatch(r"[+-]?\d+(?:[\.,]\d+)?(?:[eE][+-]?\d+)?", normalized):
            try:
                number = float(normalized.replace(",", "."))
                rounded = f"{number:.3f}"
                compact = rounded.rstrip("0").rstrip(".")
                if compact in {"", "-0"}:
                    return "0"
                return compact
            except Exception:
                return text
        return text

    def _create_docx_table(headers: List[str], rows: List[List[str]]):
        col_count = max(len(headers), max((len(r) for r in rows), default=0))
        if col_count <= 0:
            return None

        table = doc.add_table(rows=0, cols=col_count)
        table.style = "Table Grid"

        if headers:
            header_cells = table.add_row().cells
            for ci in range(col_count):
                header_cells[ci].text = headers[ci] if ci < len(headers) else ""

        for row in rows:
            row_cells = table.add_row().cells
            for ci in range(col_count):
                row_cells[ci].text = row[ci] if ci < len(row) else ""

        return table

    def _build_tables_from_payload(table_payload: Dict[str, Any]) -> List[Dict[str, Any]]:
        raw_headers = [str(h) for h in (table_payload.get("headers") or [])]
        raw_rows = [[str(cell) for cell in row] for row in (table_payload.get("rows") or [])]

        col_count = max(len(raw_headers), max((len(r) for r in raw_rows), default=0))
        if col_count <= 0:
            return []

        # Široké tabulky: rozdělit do více bloků, vždy zachovat 1. sloupec.
        needs_split = col_count > MAX_TABLE_COLUMNS_PER_BLOCK
        round_numeric = True

        headers = [_format_cell_value(h, round_numeric) for h in raw_headers]
        rows = [[_format_cell_value(cell, round_numeric) for cell in row] for row in raw_rows]

        if not needs_split:
            table_obj = _create_docx_table(headers, rows)
            return [{"table": table_obj, "part_index": 1, "part_count": 1}] if table_obj is not None else []

        first_col_idx = 0
        other_cols = list(range(1, col_count))
        chunk_size = max(1, MAX_TABLE_COLUMNS_PER_BLOCK - 1)
        col_chunks = [other_cols[i : i + chunk_size] for i in range(0, len(other_cols), chunk_size)]

        parts: List[Dict[str, Any]] = []
        for part_index, chunk in enumerate(col_chunks, start=1):
            selected = [first_col_idx] + chunk

            part_headers = [headers[idx] if idx < len(headers) else "" for idx in selected]
            part_rows = []
            for row in rows:
                part_rows.append([row[idx] if idx < len(row) else "" for idx in selected])

            table_obj = _create_docx_table(part_headers, part_rows)
            if table_obj is not None:
                parts.append({
                    "table": table_obj,
                    "part_index": part_index,
                    "part_count": len(col_chunks),
                })

        return parts

    def _add_content_after(
        paragraph,
        text: str,
        images: List,
        image_ids: List[str],
        tables: List[Dict[str, Any]] | None = None,
        one_image_per_page: bool = False,
        show_image_caption: bool = True,
        equation_mode: bool = False,
    ):
        parent = paragraph._element.getparent()
        index = parent.index(paragraph._element)

        raw_text = text or ""
        equation_blocks: List[Dict[str, str]] = []
        if equation_mode:
            equation_blocks = _build_calculation_equations(raw_text)

        text = _sanitize_section_text(text)

        if text and (not equation_mode or (equation_mode and not equation_blocks)):
            new_p = doc.add_paragraph()
            new_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = new_p.add_run(text)
            run.font.name = "Calibri"
            run.font.size = Pt(12)
            parent.insert(index + 1, new_p._element)
            index += 1

        if equation_mode:
            if equation_blocks:
                for block in equation_blocks:
                    title_p = doc.add_paragraph(block["title"])
                    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    if title_p.runs:
                        title_p.runs[0].bold = True
                        title_p.runs[0].font.name = "Calibri"
                        title_p.runs[0].font.size = Pt(12)
                    parent.insert(index + 1, title_p._element)
                    index += 1

                    for line_key in ("general", "substitution", "result"):
                        latex_line = block[line_key]
                        eq_image = _render_equation_image(latex_line)
                        if eq_image is not None:
                            eq_p = doc.add_paragraph()
                            eq_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            eq_run = eq_p.add_run()
                            eq_run.add_picture(eq_image, width=Pt(360))
                        else:
                            eq_p = doc.add_paragraph(latex_line)
                            eq_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            if eq_p.runs:
                                eq_p.runs[0].font.name = "Cambria Math"
                                eq_p.runs[0].font.size = Pt(12)

                        parent.insert(index + 1, eq_p._element)
                        index += 1

                    spacer = doc.add_paragraph("")
                    parent.insert(index + 1, spacer._element)
                    index += 1

        for table_idx, table_payload in enumerate(tables or [], start=1):
            source_file = str(table_payload.get("source_file") or "").strip()
            sheet_name = str(table_payload.get("sheet_name") or "").strip()
            caption_text = f"Tabulka {table_idx}"
            if source_file or sheet_name:
                caption_text += f" ({source_file}{' / ' + sheet_name if sheet_name else ''})"

            table_parts = _build_tables_from_payload(table_payload)
            for part in table_parts:
                part_caption_text = caption_text
                if part["part_count"] > 1:
                    part_caption_text += f" – část {part['part_index']}/{part['part_count']}"

                table_caption = doc.add_paragraph(part_caption_text)
                table_caption.alignment = WD_ALIGN_PARAGRAPH.LEFT
                if table_caption.runs:
                    table_caption.runs[0].bold = True
                    table_caption.runs[0].font.name = "Calibri"
                    table_caption.runs[0].font.size = Pt(11)
                parent.insert(index + 1, table_caption._element)
                index += 1

                table_obj = part["table"]
                parent.insert(index + 1, table_obj._element)
                index += 1

        for i, img in enumerate(images):
            if one_image_per_page and i > 0:
                break_p = doc.add_paragraph()
                break_p.add_run().add_break(WD_BREAK.PAGE)
                parent.insert(index + 1, break_p._element)
                index += 1

            img_stream = io.BytesIO()
            img.save(img_stream, format="PNG")
            img_stream.seek(0)

            img_p = doc.add_paragraph()
            img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            img_run = img_p.add_run()
            img_run.add_picture(img_stream, width=Pt(400))
            parent.insert(index + 1, img_p._element)
            index += 1

            if show_image_caption:
                image_id = image_ids[i] if i < len(image_ids) else "N/A"
                image_meta = image_registry.get(image_id)
                filename = image_meta.filename if image_meta else "neznámý soubor"
                caption = doc.add_paragraph(f"Obrázek {image_id}: {filename}")
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap_run = caption.runs[0]
                cap_run.font.name = "Calibri"
                cap_run.font.size = Pt(10)
                parent.insert(index + 1, caption._element)
                index += 1

    for para in list(doc.paragraphs):
        cleaned = para.text.strip()
        if cleaned in sections_map:
            sec = sections_map[cleaned]
            _add_content_after(
                para,
                sec["text"],
                sec["images"],
                sec["image_ids"],
                sec.get("tables", []),
                bool(sec.get("one_image_per_page", False)),
                bool(sec.get("show_image_caption", True)),
                bool(sec.get("equation_mode", False)),
            )

    bio = io.BytesIO()
    doc.save(bio)
    return bio
