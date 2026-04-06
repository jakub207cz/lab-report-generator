from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document

from pipeline.ingestion import IngestionPipeline


def _write_docx(path: Path, text: str) -> None:
    doc = Document()
    doc.add_paragraph(text)
    doc.save(path)


def _write_xlsx(path: Path) -> None:
    df = pd.DataFrame(
        {
            "Napeti (V)": [5.0, 12.0],
            "Proud (A)": [0.5, 0.8],
        }
    )
    df.to_excel(path, index=False)


def test_ingest_single_txt_file_path_to_normalized_result(tmp_path: Path) -> None:
    txt_path = tmp_path / "zadani.txt"
    txt_path.write_text("Mereni stabilizovaneho zdroje", encoding="utf-8")

    pipeline = IngestionPipeline(enable_ocr=False)
    result = pipeline.ingest_file_path(txt_path, section_hint="assignment")

    assert result.metadata["source_count"] == 1
    assert result.metadata["chunk_count"] >= 1
    assert any("stabilizovaneho zdroje" in chunk.text for chunk in result.chunks)
    assert all(chunk.section_hint == "assignment" for chunk in result.chunks)


def test_ingest_docx_and_xlsx_paths_to_normalized_result(tmp_path: Path) -> None:
    docx_path = tmp_path / "teorie.docx"
    xlsx_path = tmp_path / "data.xlsx"

    _write_docx(docx_path, "Ohmuv zakon popisuje zavislost mezi U, I, R.")
    _write_xlsx(xlsx_path)

    pipeline = IngestionPipeline(enable_ocr=False)
    result = pipeline.ingest_file_paths([docx_path, xlsx_path], section_hint="theory")

    assert result.metadata["source_count"] == 2
    assert result.metadata["table_count"] >= 1
    assert any("Ohmuv zakon" in chunk.text for chunk in result.chunks)

    table = result.tables[0]
    assert "Napeti (V)" in table.headers
    assert table.units is not None
    assert table.units[0] == "V"
    assert table.section_hint == "theory"


def test_ingest_csv_file_path_creates_text_and_table(tmp_path: Path) -> None:
    csv_path = tmp_path / "mereni.csv"
    csv_path.write_text("Napeti (V),Proud (A)\n5,0.5\n12,0.8\n", encoding="utf-8")

    pipeline = IngestionPipeline(enable_ocr=False)
    result = pipeline.ingest_file_path(csv_path, section_hint="data")

    assert any("Napeti (V),Proud (A)" in chunk.text for chunk in result.chunks)
    assert len(result.tables) == 1
    assert result.tables[0].headers == ["Napeti (V)", "Proud (A)"]
    assert result.tables[0].section_hint == "data"


def test_extract_assignment_sections_from_text_map() -> None:
    sample_assignment_text = """
    Zadání
    Ověřte vlastnosti feromagnetického materiálu.

    Teoretický úvod
    Popište hysterezní smyčku a magnetizaci.

    Schéma zapojení
    Obrázek schématu je součástí zadání.

    Postup měření
    Připravte měřicí pracoviště a změřte body křivky.

    Příklad tabulek
    Uveďte tabulku H a B.

    Příklad výpočtu
    Vypočtěte relativní permeabilitu.

    Předpokládaný průběh grafů
    Náčrt B=f(H) a hysterezní smyčky.

    Závěr
    Zhodnoťte výsledky měření a odchylky.
    """

    pipeline = IngestionPipeline(enable_ocr=False)
    sections = pipeline._extract_assignment_sections(sample_assignment_text)

    expected_keys = {
        "assignment",
        "theory",
        "schema",
        "procedure",
        "tables_example",
        "calculation_example",
        "expected_graphs",
        "conclusion",
    }

    assert set(sections.keys()) == expected_keys
    assert "hysterezní smyčku" in sections["theory"]["text"].lower()
    assert "permeabilitu" in sections["calculation_example"]["text"].lower()
    assert sections["conclusion"]["title"] == "Závěr"


def test_ingest_real_assignment_pdf_file() -> None:
    project_root = Path(__file__).resolve().parents[1]
    pdf_path = project_root / "11_Feromag_mat IoT.pdf"

    assert pdf_path.exists()

    pipeline = IngestionPipeline(enable_ocr=False)
    result = pipeline.ingest_file_path(pdf_path, section_hint="assignment")

    assert result.metadata["source_count"] == 1
    assert result.metadata["chunk_count"] >= 1
    assert result.metadata["figure_count"] >= 1

    assignment_sections = result.metadata.get("assignment_sections")
    assert isinstance(assignment_sections, dict)
    assert len(assignment_sections) >= 4
    assert {"theory", "schema", "procedure", "conclusion"}.intersection(assignment_sections.keys())

    saved_figures = result.metadata.get("saved_figures")
    assert isinstance(saved_figures, list)
    assert len(saved_figures) >= 1
    for saved in saved_figures:
        assert Path(saved["path"]).exists()


def test_ingest_real_xlsx_extracts_tables_and_graphs() -> None:
    project_root = Path(__file__).resolve().parents[1]
    xlsx_path = project_root / "em-elab-11.xlsx"

    assert xlsx_path.exists()

    pipeline = IngestionPipeline(enable_ocr=False)
    result = pipeline.ingest_file_path(xlsx_path, section_hint="data")

    assert result.metadata["source_count"] == 1
    assert result.metadata["table_count"] >= 2
    assert result.metadata["figure_count"] >= 1

    table_layouts = result.metadata.get("table_layouts")
    assert isinstance(table_layouts, list)
    assert len(table_layouts) == result.metadata["table_count"]
    assert all(layout.get("row_count", 0) >= 2 for layout in table_layouts)
    assert all(layout.get("col_count", 0) >= 2 for layout in table_layouts)
    assert all("!" in str(layout.get("sheet_range", "")) for layout in table_layouts)

    flat_table_cells = [str(cell).lower() for table in result.tables for row in table.rows for cell in row]
    flat_headers = [str(header).lower() for table in result.tables for header in table.headers]
    combined = flat_table_cells + flat_headers
    assert any("pztr" in cell or "pfe" in cell or "b[t]" in cell or "b [t]" in cell for cell in combined)

    assert any(chunk.type == "table" for chunk in result.chunks)
    assert any(chunk.type == "figure" for chunk in result.chunks)
    assert all(fig.source_file == "em-elab-11.xlsx" for fig in result.figures)

    saved_figures = result.metadata.get("saved_figures")
    assert isinstance(saved_figures, list)
    assert len(saved_figures) >= 1
    for saved in saved_figures:
        assert Path(saved["path"]).exists()
