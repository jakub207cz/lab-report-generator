from __future__ import annotations

import io
import json
from pathlib import Path

from docx import Document
from PIL import Image

from pipeline.schemas import ImageAsset, LabReportData
from render.docx_writer import fill_template_docx


def _all_cell_texts(doc: Document) -> list[str]:
    texts: list[str] = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = (cell.text or "").strip()
                if text:
                    texts.append(text)
    return texts


def test_fill_template_docx_populates_cover_and_sections() -> None:
    project_root = Path(__file__).resolve().parents[1]
    template_path = project_root / "DOCS" / "sablona_elab_em.docx"

    ai_data = LabReportData(
        teorie="Teorie test",
        postup="Postup test",
        priklad_vypoctu="Vypocet test",
        zaver="Zaver test",
        image_references=[],
    )

    schema_img = Image.new("RGB", (64, 64), color="white")
    wave_img = Image.new("RGB", (64, 64), color="black")

    inputs_map = {
        "assignment_text": "Toto je testovaci zadani pro overeni zapisu do desek.",
        "instruments_text": "Multimetr, osciloskop",
        "waveforms_text": "Prubeh je sinusovy.",
        "schema_images": [schema_img],
        "schema_image_ids": ["IMG-001"],
        "waveforms_images": [wave_img],
        "waveforms_image_ids": ["IMG-002"],
        "data_images": [],
        "data_image_ids": [],
        "username": "Jan Novak",
        "topic": "Mereni feromagnetickych materialu",
    }

    image_registry = {
        "IMG-001": ImageAsset(image_id="IMG-001", filename="schema.png", section="schema"),
        "IMG-002": ImageAsset(image_id="IMG-002", filename="graf.png", section="waveforms"),
    }

    out = fill_template_docx(
        str(template_path),
        inputs_map,
        ai_data,
        image_registry,
        username="Jan Novak",
        topic="Mereni feromagnetickych materialu",
    )

    generated = Document(io.BytesIO(out.getvalue()))
    cell_texts = _all_cell_texts(generated)

    assert any("Mereni feromagnetickych materialu" in text for text in cell_texts)
    assert any("Jan Novak" in text for text in cell_texts)
    assert any("Toto je testovaci zadani" in text for text in cell_texts)

    header_texts = [
        p.text.strip()
        for section in generated.sections
        for p in section.header.paragraphs
        if p.text.strip()
    ]
    assert any("Jan Novak" in text for text in header_texts)

    header_xml_chunks = [
        p._p.xml
        for section in generated.sections
        for p in section.header.paragraphs
    ]
    assert any("PAGE" in xml for xml in header_xml_chunks)
    assert any('w:val="right"' in xml for xml in header_xml_chunks)


def test_fill_template_docx_cover_assignment_uses_clean_numbered_items() -> None:
    project_root = Path(__file__).resolve().parents[1]
    template_path = project_root / "DOCS" / "sablona_elab_em.docx"

    ai_data = LabReportData(
        teorie="",
        postup="",
        priklad_vypoctu="",
        zaver="",
        image_references=[],
    )

    inputs_map = {
        "assignment_text": """
--- 11_Feromag_mat IoT.pdf ---
FEROMAG_MAT.DOC
1
Úloha č. ...........
Měření feromagnetických materiálů
Zadání:
1. První bod
s pokračováním na dalším řádku.
2. Druhý bod zadání.
Teoretický úvod:
Toto už se na desky nemá propsat.
""",
        "instruments_text": "",
        "waveforms_text": "",
        "schema_images": [],
        "schema_image_ids": [],
        "waveforms_images": [],
        "waveforms_image_ids": [],
        "data_images": [],
        "data_image_ids": [],
        "username": "Jan Novak",
        "topic": "Mereni",
    }

    out = fill_template_docx(
        str(template_path),
        inputs_map,
        ai_data,
        image_registry={},
        username="Jan Novak",
        topic="Mereni",
    )

    generated = Document(io.BytesIO(out.getvalue()))
    cell_texts = _all_cell_texts(generated)

    assert any("1. První bod s pokračováním na dalším řádku." in text for text in cell_texts)
    assert any("2. Druhý bod zadání." in text for text in cell_texts)
    assert not any("--- 11_Feromag_mat IoT.pdf ---" in text for text in cell_texts)
    assert not any("FEROMAG_MAT.DOC" in text for text in cell_texts)
    assert not any("Teoretický úvod" in text for text in cell_texts)


def test_fill_template_docx_splits_wide_table_and_keeps_first_column() -> None:
    project_root = Path(__file__).resolve().parents[1]
    template_path = project_root / "DOCS" / "sablona_elab_em.docx"

    ai_data = LabReportData(
        teorie="",
        postup="",
        priklad_vypoctu="",
        zaver="",
        image_references=[],
    )

    wide_headers = ["Udaj"] + [f"M{i}" for i in range(1, 13)]
    wide_rows = [
        ["B [T]"] + [str(v) for v in [0.12345, 0.23456, 0.34567, 0.45678, 0.56789, 0.67891, 0.78912, 0.89123, 1.23456, 2.34567, 3.45678, 4.56789]],
        ["Pztr [W/kg]"] + [str(v) for v in [1.11119, 1.22229, 1.33339, 1.44449, 1.55559, 1.66669, 1.77779, 1.88889, 1.99999, 2.11119, 2.22229, 2.33339]],
    ]

    inputs_map = {
        "assignment_text": "",
        "instruments_text": "",
        "data_text": "Tabulka z měření:",
        "data_tables": [
            {
                "source_file": "test.xlsx",
                "sheet_name": "Sheet1!A1:M3",
                "headers": wide_headers,
                "rows": wide_rows,
            }
        ],
        "waveforms_text": "",
        "schema_images": [],
        "schema_image_ids": [],
        "waveforms_images": [],
        "waveforms_image_ids": [],
        "data_images": [],
        "data_image_ids": [],
        "username": "Tester",
        "topic": "Široká tabulka",
    }

    out = fill_template_docx(
        str(template_path),
        inputs_map,
        ai_data,
        image_registry={},
        username="Tester",
        topic="Široká tabulka",
    )

    generated = Document(io.BytesIO(out.getvalue()))

    # 1. tabulka je obálka; široká tabulka se má rozdělit minimálně na 2 části
    assert len(generated.tables) >= 3

    first_part = generated.tables[1]
    second_part = generated.tables[2]

    assert first_part.cell(0, 0).text.strip() == "Udaj"
    assert second_part.cell(0, 0).text.strip() == "Udaj"

    # První sloupec s názvem řádku musí být zachovaný v každé části
    assert first_part.cell(1, 0).text.strip() == "B [T]"
    assert second_part.cell(1, 0).text.strip() == "B [T]"

    # Číselné hodnoty u široké tabulky se mají zaokrouhlit na 3 desetinná místa
    table_texts = [
        (cell.text or "").strip()
        for table in generated.tables[1:]
        for row in table.rows
        for cell in row.cells
    ]
    assert "0.123" in table_texts
    assert "1.111" in table_texts

    para_texts = [(p.text or "").strip() for p in generated.paragraphs]
    assert any("část 1/" in t for t in para_texts)
    assert any("část 2/" in t for t in para_texts)


def test_fill_template_docx_rounds_all_numeric_values_and_zero_compacts() -> None:
    project_root = Path(__file__).resolve().parents[1]
    template_path = project_root / "DOCS" / "sablona_elab_em.docx"

    ai_data = LabReportData(
        teorie="",
        postup="",
        priklad_vypoctu="",
        zaver="",
        image_references=[],
    )

    inputs_map = {
        "assignment_text": "",
        "instruments_text": "",
        "data_text": "Tabulka z měření:",
        "data_tables": [
            {
                "source_file": "test.xlsx",
                "sheet_name": "Sheet1",
                "headers": ["Veličina", "Hodnota"],
                "rows": [
                    ["I", "12.34567"],
                    ["A", "27.700"],
                    ["B", "0.600"],
                    ["C", "2.500"],
                    ["D", "0.060"],
                    ["U", "0.000"],
                ],
            }
        ],
        "waveforms_text": "",
        "schema_images": [],
        "schema_image_ids": [],
        "waveforms_images": [],
        "waveforms_image_ids": [],
        "data_images": [],
        "data_image_ids": [],
        "username": "Tester",
        "topic": "Zaokrouhlení tabulek",
    }

    out = fill_template_docx(
        str(template_path),
        inputs_map,
        ai_data,
        image_registry={},
        username="Tester",
        topic="Zaokrouhlení tabulek",
    )

    generated = Document(io.BytesIO(out.getvalue()))

    # 1. tabulka = obálka, 2. tabulka = data
    values = [
        (cell.text or "").strip()
        for row in generated.tables[1].rows
        for cell in row.cells
    ]
    assert "12.346" in values
    assert "27.7" in values
    assert "0.6" in values
    assert "2.5" in values
    assert "0.06" in values
    assert "0" in values
    assert "27.700" not in values
    assert "0.600" not in values
    assert "2.500" not in values
    assert "0.060" not in values
    assert "0.000" not in values


def test_fill_template_docx_graphs_have_page_break_between_images() -> None:
    project_root = Path(__file__).resolve().parents[1]
    template_path = project_root / "DOCS" / "sablona_elab_em.docx"

    ai_data = LabReportData(teorie="", postup="", priklad_vypoctu="", zaver="", image_references=[])

    inputs_map = {
        "assignment_text": "",
        "instruments_text": "",
        "data_text": "",
        "data_tables": [],
        "waveforms_text": "",  # text část necháváme prázdnou, testujeme jen stránkování obrázků
        "schema_images": [],
        "schema_image_ids": [],
        "waveforms_images": [
            Image.new("RGB", (64, 64), color="red"),
            Image.new("RGB", (64, 64), color="green"),
            Image.new("RGB", (64, 64), color="blue"),
        ],
        "waveforms_image_ids": ["IMG-001", "IMG-002", "IMG-003"],
        "data_images": [],
        "data_image_ids": [],
        "username": "Tester",
        "topic": "Grafy po jedné stránce",
    }

    image_registry = {
        "IMG-001": ImageAsset(image_id="IMG-001", filename="g1.png", section="waveforms"),
        "IMG-002": ImageAsset(image_id="IMG-002", filename="g2.png", section="waveforms"),
        "IMG-003": ImageAsset(image_id="IMG-003", filename="g3.png", section="waveforms"),
    }

    out = fill_template_docx(
        str(template_path),
        inputs_map,
        ai_data,
        image_registry=image_registry,
        username="Tester",
        topic="Grafy po jedné stránce",
    )

    generated = Document(io.BytesIO(out.getvalue()))
    page_break_count = sum(p._p.xml.count('w:type="page"') for p in generated.paragraphs)

    # Pro 3 grafy očekáváme minimálně 2 page breaky mezi nimi.
    assert page_break_count >= 2


def test_fill_template_docx_graphs_are_inserted_without_caption() -> None:
    project_root = Path(__file__).resolve().parents[1]
    template_path = project_root / "DOCS" / "sablona_elab_em.docx"

    ai_data = LabReportData(teorie="", postup="", priklad_vypoctu="", zaver="", image_references=[])
    inputs_map = {
        "assignment_text": "",
        "instruments_text": "",
        "data_text": "",
        "data_tables": [],
        "waveforms_text": "",
        "schema_images": [],
        "schema_image_ids": [],
        "waveforms_images": [Image.new("RGB", (64, 64), color="purple")],
        "waveforms_image_ids": ["IMG-123"],
        "data_images": [],
        "data_image_ids": [],
        "username": "Tester",
        "topic": "Graf bez titulku",
    }
    image_registry = {
        "IMG-123": ImageAsset(image_id="IMG-123", filename="graf.png", section="waveforms"),
    }

    out = fill_template_docx(
        str(template_path),
        inputs_map,
        ai_data,
        image_registry=image_registry,
        username="Tester",
        topic="Graf bez titulku",
    )

    generated = Document(io.BytesIO(out.getvalue()))
    paragraph_texts = [(p.text or "").strip() for p in generated.paragraphs if (p.text or "").strip()]

    assert not any("Obrázek IMG-123" in text for text in paragraph_texts)


def test_fill_template_docx_renders_calculations_from_llm_json() -> None:
    project_root = Path(__file__).resolve().parents[1]
    template_path = project_root / "DOCS" / "sablona_elab_em.docx"

    calc_payload = {
        "items": [
            {
                "title": "Intenzita magnetického pole Hm",
                "general_formula_latex": r"H_M = \\frac{\\sqrt{2}\\cdot I_1 \\cdot N_1}{l_{str}}",
                "substitution_formula_latex": r"H_M = \\frac{\\sqrt{2}\\cdot 0.2 \\cdot 100}{0.15}",
                "compute_expression": "(sqrt(2)*0.2*100)/0.15",
                "result_symbol_latex": "H_M",
                "result_unit_latex": "A/m",
                "variables": {},
            }
        ]
    }

    ai_data = LabReportData(
        teorie="",
        postup="",
        priklad_vypoctu=json.dumps(calc_payload, ensure_ascii=False),
        zaver="",
        image_references=[],
    )

    inputs_map = {
        "assignment_text": "",
        "instruments_text": "",
        "data_text": "",
        "data_tables": [],
        "waveforms_text": "",
        "schema_images": [],
        "schema_image_ids": [],
        "waveforms_images": [],
        "waveforms_image_ids": [],
        "data_images": [],
        "data_image_ids": [],
        "username": "Tester",
        "topic": "Rovnice z LLM JSON",
        "report_scope": "full",
    }

    out = fill_template_docx(
        str(template_path),
        inputs_map,
        ai_data,
        image_registry={},
        username="Tester",
        topic="Rovnice z LLM JSON",
    )

    generated = Document(io.BytesIO(out.getvalue()))
    paragraph_texts = [(p.text or "").strip() for p in generated.paragraphs if (p.text or "").strip()]

    assert any("Intenzita magnetického pole Hm" in text for text in paragraph_texts)
    assert not any('"items"' in text for text in paragraph_texts)
