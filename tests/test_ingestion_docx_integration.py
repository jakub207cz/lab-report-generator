from __future__ import annotations

import io
from pathlib import Path

from docx import Document
from PIL import Image

from pipeline.schemas import LabReportData
from pipeline.ingestion import IngestionPipeline
from render.docx_writer import fill_template_docx


def _paragraph_texts(doc: Document) -> list[str]:
    return [p.text.strip() for p in doc.paragraphs if (p.text or "").strip()]


def test_xlsx_ingestion_to_docx_preview_output() -> None:
    project_root = Path(__file__).resolve().parents[1]
    template_path = project_root / "DOCS" / "sablona_elab_em.docx"
    xlsx_path = project_root / "em-elab-11.xlsx"

    assert template_path.exists()
    assert xlsx_path.exists()

    ingestion = IngestionPipeline(enable_ocr=False)
    ingested = ingestion.ingest_file_path(xlsx_path, section_hint="data")

    assert len(ingested.tables) >= 1
    assert len(ingested.figures) >= 1

    data_tables = [
        {
            "source_file": t.source_file,
            "sheet_name": t.sheet_name,
            "headers": [str(h) for h in t.headers],
            "rows": [[str(cell) for cell in row] for row in t.rows],
        }
        for t in ingested.tables
    ]
    figure_lines = [
        f"- {fig.figure_id}: {(fig.ocr_text or '').strip()}".strip()
        for fig in ingested.figures
    ]

    ai_data = LabReportData(
        teorie="Teoretický úvod test integrace XLSX -> DOCX.",
        postup="Postup měření test integrace XLSX -> DOCX.",
        priklad_vypoctu="Příklad výpočtu test integrace XLSX -> DOCX.",
        zaver="Závěr test integrace XLSX -> DOCX.",
        image_references=[fig.figure_id for fig in ingested.figures],
    )

    wave_images = []
    wave_image_ids: list[str] = []
    saved_figures = ingested.metadata.get("saved_figures", []) if isinstance(ingested.metadata, dict) else []
    for idx, saved in enumerate(saved_figures[:4], start=1):
        image_path = Path(str(saved.get("path", "")))
        if not image_path.exists():
            continue
        try:
            wave_images.append(Image.open(image_path).convert("RGB"))
        except Exception:
            continue
        wave_image_ids.append(f"IMG-XLSX-{idx:03d}")

    inputs_map = {
        "assignment_text": "Integrační test: napojení ingestion a docx writeru.",
        "instruments_text": "Multimetr, wattmetr, osciloskop",
        "data_text": "Tabulky z XLSX byly vloženy níže jako skutečné tabulky.",
        "data_tables": data_tables,
        "waveforms_text": "Detekované grafy z XLSX:\n" + "\n".join(figure_lines),
        "schema_images": [],
        "schema_image_ids": [],
    "waveforms_images": wave_images,
    "waveforms_image_ids": wave_image_ids,
        "data_images": [],
        "data_image_ids": [],
        "username": "Test Uživatel",
        "topic": "Měření feromagnetických materiálů",
    }

    out = fill_template_docx(
        template_path=str(template_path),
        inputs_map=inputs_map,
        ai_content=ai_data,
        image_registry={},
        username="Test Uživatel",
        topic="Měření feromagnetických materiálů",
    )

    generated_bytes = out.getvalue()
    assert len(generated_bytes) > 0

    output_dir = project_root / "tests" / "output"
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / "preview_em_elab_11_from_xlsx.docx"
    output_path.write_bytes(generated_bytes)
    assert output_path.exists()
    assert output_path.stat().st_size > 0

    generated_doc = Document(io.BytesIO(generated_bytes))
    paragraph_texts = _paragraph_texts(generated_doc)

    assert any("Detekované grafy z XLSX" in txt for txt in paragraph_texts)
    assert any("Chartsheet" in txt or "Worksheet chart" in txt for txt in paragraph_texts)
    assert any("skutečné tabulky" in txt for txt in paragraph_texts)
    assert any("Příklad výpočtu test integrace" in txt for txt in paragraph_texts)
    assert len(generated_doc.tables) >= 1 + len(data_tables)

    all_table_cells = [
        (cell.text or "").strip()
        for table in generated_doc.tables[1:]
        for row in table.rows
        for cell in row.cells
    ]
    assert any("N1" in txt for txt in all_table_cells)
    assert any("Epsteinův přístroj" in txt for txt in all_table_cells)
    if wave_images:
        assert len(generated_doc.inline_shapes) >= len(wave_images)
