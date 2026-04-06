from __future__ import annotations

import io
import os
from typing import Dict, List, Literal, Optional, Tuple

import pandas as pd
import pypdf
from docx import Document
from dotenv import load_dotenv
from fastapi import FastAPI, File, Form, UploadFile
from fastapi.responses import FileResponse, JSONResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from PIL import Image, UnidentifiedImageError
import uvicorn

from pipeline.generation import generate_lab_report as generate_lab_report_structured
from pipeline.ingestion import BinarySource, IngestionPipeline
from pipeline.schemas import ImageAsset
from pipeline.validation import run_quality_check
from render.docx_writer import fill_template_docx as fill_template_docx_structured

ReportScope = Literal["full", "preparation", "ending"]

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
load_dotenv(dotenv_path=os.path.join(BASE_DIR, ".env"), override=False)

WEB_DIR = os.path.join(BASE_DIR, "web")
app = FastAPI(title="Lab Report Generator API", version="1.0.0")
app.mount("/static", StaticFiles(directory=WEB_DIR), name="static")


def _safe_decode(raw: bytes) -> str:
    for enc in ("utf-8", "cp1250", "latin-1"):
        try:
            return raw.decode(enc)
        except Exception:
            continue
    return ""


async def extract_content_from_uploadfiles(
    files: Optional[List[UploadFile]],
    section_name: str,
    image_counter_start: int,
) -> Tuple[str, List[Image.Image], List[ImageAsset], List[Dict[str, object]], int]:
    text_content = ""
    images: List[Image.Image] = []
    image_assets: List[ImageAsset] = []
    tables_payload: List[Dict[str, object]] = []
    image_counter = image_counter_start

    if not files:
        return text_content, images, image_assets, tables_payload, image_counter

    spreadsheet_sources: List[BinarySource] = []
    spreadsheet_raw_by_name: Dict[str, bytes] = {}

    for uploaded_file in files:
        try:
            raw = await uploaded_file.read()
            content_type = uploaded_file.content_type or ""
            filename = uploaded_file.filename or "unknown"

            if content_type.startswith("image/") or filename.lower().endswith((".png", ".jpg", ".jpeg", ".bmp", ".webp")):
                image = Image.open(io.BytesIO(raw)).convert("RGB")
                images.append(image)

                image_id = f"IMG-{image_counter:03d}"
                image_counter += 1
                image_assets.append(ImageAsset(image_id=image_id, filename=filename, section=section_name))

            elif filename.lower().endswith((".xlsx", ".xls")) or "spreadsheet" in content_type:
                spreadsheet_sources.append(
                    BinarySource(
                        filename=filename,
                        data=raw,
                        section_hint=section_name,
                        mime_type=content_type,
                    )
                )
                spreadsheet_raw_by_name[filename] = raw
                continue

            elif filename.lower().endswith(".docx") or "wordprocessing" in content_type:
                doc = Document(io.BytesIO(raw))
                full_text = [p.text for p in doc.paragraphs if p.text]
                text_content += f"\n--- {filename} ---\n" + "\n".join(full_text) + "\n"

            elif filename.lower().endswith(".pdf") or "pdf" in content_type:
                reader = pypdf.PdfReader(io.BytesIO(raw))
                pdf_text = ""
                for page in reader.pages:
                    pdf_text += (page.extract_text() or "") + "\n"
                text_content += f"\n--- {filename} ---\n" + pdf_text + "\n"

            elif filename.lower().endswith((".txt", ".csv")) or content_type == "text/plain":
                text_content += f"\n--- {filename} ---\n" + _safe_decode(raw) + "\n"

        except UnidentifiedImageError:
            continue
        except Exception:
            continue

    if spreadsheet_sources:
        parsed_tables_by_source: Dict[str, List[object]] = {}
        saved_figures: List[Dict[str, object]] = []

        try:
            ingestion = IngestionPipeline(enable_ocr=False)
            parsed = ingestion.ingest_sources(spreadsheet_sources)

            for table in parsed.tables:
                parsed_tables_by_source.setdefault(table.source_file, []).append(table)

            if isinstance(parsed.metadata, dict):
                saved_figures = parsed.metadata.get("saved_figures", []) or []
        except Exception:
            parsed_tables_by_source = {}
            saved_figures = []

        for source in spreadsheet_sources:
            filename = source.filename
            file_tables = parsed_tables_by_source.get(filename, [])

            if file_tables:
                for table in file_tables:
                    headers = [str(h) for h in table.headers]
                    rows = [[str(cell) for cell in row] for row in table.rows]
                    tables_payload.append(
                        {
                            "source_file": filename,
                            "sheet_name": table.sheet_name,
                            "headers": headers,
                            "rows": rows,
                        }
                    )
            else:
                try:
                    df = pd.read_excel(io.BytesIO(spreadsheet_raw_by_name[filename]))
                    text_content += f"\n--- {filename} ---\n" + df.to_markdown(index=False) + "\n"
                except Exception:
                    pass

        # Přidej všechny grafy z všech XLSX souborů.
        for saved in saved_figures:
            image_path = str(saved.get("path", "")).strip()
            if not image_path:
                continue
            try:
                image = Image.open(image_path).convert("RGB")
            except Exception:
                continue

            image_id = f"IMG-{image_counter:03d}"
            image_counter += 1
            images.append(image)
            image_assets.append(ImageAsset(image_id=image_id, filename=os.path.basename(image_path), section=section_name))

    return text_content, images, image_assets, tables_payload, image_counter


async def extract_assignment_fallback_images(
    assignment_files: Optional[List[UploadFile]],
    image_counter_start: int,
) -> Tuple[List[Image.Image], List[ImageAsset], List[Image.Image], List[ImageAsset], int]:
    schema_images: List[Image.Image] = []
    schema_assets: List[ImageAsset] = []
    waveforms_images: List[Image.Image] = []
    waveforms_assets: List[ImageAsset] = []
    image_counter = image_counter_start

    if not assignment_files:
        return schema_images, schema_assets, waveforms_images, waveforms_assets, image_counter

    sources: List[BinarySource] = []
    for uploaded_file in assignment_files:
        try:
            raw = await uploaded_file.read()
            await uploaded_file.seek(0)
        except Exception:
            continue

        filename = uploaded_file.filename or "unknown"
        content_type = uploaded_file.content_type or ""
        if not (filename.lower().endswith(".pdf") or "pdf" in content_type):
            continue

        sources.append(
            BinarySource(
                filename=filename,
                data=raw,
                section_hint="assignment",
                mime_type=content_type,
            )
        )

    if not sources:
        return schema_images, schema_assets, waveforms_images, waveforms_assets, image_counter

    ingestion = IngestionPipeline(enable_ocr=False)
    result = ingestion.ingest_sources(sources)

    saved_figures = result.metadata.get("saved_figures", [])
    section_by_figure_id = {
        figure.figure_id: (figure.section_hint or "")
        for figure in result.figures
    }

    for saved_figure in saved_figures:
        figure_id = saved_figure.get("figure_id")
        image_path = saved_figure.get("path")
        if not image_path:
            continue

        section_hint = section_by_figure_id.get(figure_id, "")
        if not (
            section_hint.startswith("assignment:schema")
            or section_hint.startswith("assignment:expected_graphs")
        ):
            continue

        try:
            extracted_img = Image.open(image_path).convert("RGB")
        except Exception:
            continue

        image_id = f"IMG-{image_counter:03d}"
        image_counter += 1
        asset = ImageAsset(
            image_id=image_id,
            filename=os.path.basename(image_path),
            section="schema" if section_hint.startswith("assignment:schema") else "waveforms",
        )

        if section_hint.startswith("assignment:schema"):
            schema_images.append(extracted_img)
            schema_assets.append(asset)
        else:
            waveforms_images.append(extracted_img)
            waveforms_assets.append(asset)

    return schema_images, schema_assets, waveforms_images, waveforms_assets, image_counter


async def extract_assignment_sections_from_files(assignment_files: Optional[List[UploadFile]]) -> Dict[str, Dict[str, str]]:
    if not assignment_files:
        return {}

    sources: List[BinarySource] = []
    for uploaded_file in assignment_files:
        try:
            raw = await uploaded_file.read()
            await uploaded_file.seek(0)
        except Exception:
            continue

        filename = uploaded_file.filename or "unknown"
        content_type = uploaded_file.content_type or ""
        if not (filename.lower().endswith(".pdf") or "pdf" in content_type):
            continue

        sources.append(
            BinarySource(
                filename=filename,
                data=raw,
                section_hint="assignment",
                mime_type=content_type,
            )
        )

    if not sources:
        return {}

    ingestion = IngestionPipeline(enable_ocr=False)
    result = ingestion.ingest_sources(sources)
    sections = result.metadata.get("assignment_sections")
    return sections if isinstance(sections, dict) else {}


def _reroute_xlsx_charts_from_data_to_waveforms(
    data_images: List[Image.Image],
    data_assets: List[ImageAsset],
    waveforms_images: List[Image.Image],
    waveforms_assets: List[ImageAsset],
    data_text: str,
    waveforms_text: str,
) -> Tuple[List[Image.Image], List[ImageAsset], List[Image.Image], List[ImageAsset], str, str]:
    kept_data_images: List[Image.Image] = []
    kept_data_assets: List[ImageAsset] = []

    for img, asset in zip(data_images, data_assets):
        filename_norm = (asset.filename or "").strip().lower()
        is_xlsx_chart = filename_norm.startswith("xlsx-chart") and filename_norm.endswith(".png")

        if is_xlsx_chart:
            waveforms_images.append(img)
            waveforms_assets.append(
                ImageAsset(
                    image_id=asset.image_id,
                    filename=asset.filename,
                    section="waveforms",
                )
            )
            continue

        kept_data_images.append(img)
        kept_data_assets.append(asset)

    return kept_data_images, kept_data_assets, waveforms_images, waveforms_assets, data_text, waveforms_text


def _build_docx(
    topic: str,
    username: str,
    ai_data,
    inputs_map: Dict,
    image_registry: Dict[str, ImageAsset],
) -> io.BytesIO:
    doc = Document()
    doc.add_heading("Laboratorní protokol", level=1)
    if username.strip():
        doc.add_paragraph(f"Vypracoval: {username.strip()}")
    doc.add_paragraph(f"Téma: {topic}")

    sections = [
        ("Teoretický úvod", ai_data.teorie),
        ("Postup měření", ai_data.postup),
        ("Příklad výpočtu", ai_data.priklad_vypoctu),
        ("Závěr", ai_data.zaver),
    ]

    for title, content in sections:
        doc.add_heading(title, level=2)
        doc.add_paragraph(content or "")

    image_groups = [
        ("Grafické průběhy", inputs_map.get("waveforms_images", []), inputs_map.get("waveforms_image_ids", [])),
        ("Schéma zapojení", inputs_map.get("schema_images", []), inputs_map.get("schema_image_ids", [])),
        ("Naměřené hodnoty (obrázky)", inputs_map.get("data_images", []), inputs_map.get("data_image_ids", [])),
    ]

    for group_name, group_images, group_ids in image_groups:
        if not group_images:
            continue
        doc.add_heading(group_name, level=2)
        for i, img in enumerate(group_images):
            img_stream = io.BytesIO()
            img.save(img_stream, format="PNG")
            img_stream.seek(0)
            doc.add_picture(img_stream)
            img_id = group_ids[i] if i < len(group_ids) else "N/A"
            meta = image_registry.get(img_id)
            filename = meta.filename if meta else "neznámý soubor"
            doc.add_paragraph(f"Obrázek {img_id}: {filename}")

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out


def _resolve_report_scope(
    report_scope: str,
    generate_full_report: bool,
    generate_preparation_only: bool,
    generate_ending_only: bool,
) -> ReportScope:
    normalized_scope = (report_scope or "").strip().lower()
    if normalized_scope in ("full", "preparation", "ending"):
        return normalized_scope  # type: ignore[return-value]

    selected_count = int(bool(generate_full_report)) + int(bool(generate_preparation_only)) + int(bool(generate_ending_only))
    if selected_count == 0:
        return "full"
    if selected_count > 1:
        raise ValueError("Vyber pouze jeden režim generování.")
    if generate_preparation_only:
        return "preparation"
    if generate_ending_only:
        return "ending"
    return "full"


@app.get("/")
def index() -> FileResponse:
    return FileResponse(os.path.join(WEB_DIR, "index.html"))


@app.post("/api/generate")
async def generate_report(
    topic: str = Form(...),
    username: str = Form(""),
    is_handwritten: bool = Form(False),
    report_scope: str = Form("full"),
    generate_full_report: bool = Form(False),
    generate_preparation_only: bool = Form(False),
    generate_ending_only: bool = Form(False),
    model_name: str = Form("gemini-2.5-flash"),
    api_key: str = Form(...),
    assignment_files: Optional[List[UploadFile]] = File(None),
    instruments_files: Optional[List[UploadFile]] = File(None),
    data_files: Optional[List[UploadFile]] = File(None),
    theory_files: Optional[List[UploadFile]] = File(None),
    waveforms_files: Optional[List[UploadFile]] = File(None),
    procedure_files: Optional[List[UploadFile]] = File(None),
    conclusion_files: Optional[List[UploadFile]] = File(None),
    schema_files: Optional[List[UploadFile]] = File(None),
):
    resolved_api_key = (api_key or "").strip()
    if not resolved_api_key:
        return JSONResponse(status_code=400, content={"error": "Chybí API klíč."})

    try:
        resolved_report_scope = _resolve_report_scope(
            report_scope=report_scope,
            generate_full_report=generate_full_report,
            generate_preparation_only=generate_preparation_only,
            generate_ending_only=generate_ending_only,
        )
    except ValueError as exc:
        return JSONResponse(status_code=400, content={"error": str(exc)})

    image_counter = 1

    (
        fallback_schema_images,
        fallback_schema_assets,
        fallback_waveforms_images,
        fallback_waveforms_assets,
        image_counter,
    ) = await extract_assignment_fallback_images(assignment_files, image_counter)

    assignment_sections = await extract_assignment_sections_from_files(assignment_files)

    assignment_text, assignment_images, assignment_assets, assignment_tables, image_counter = await extract_content_from_uploadfiles(assignment_files, "assignment", image_counter)
    instruments_text, instruments_images, instruments_assets, instruments_tables, image_counter = await extract_content_from_uploadfiles(instruments_files, "instruments", image_counter)
    data_text, data_images, data_assets, data_tables, image_counter = await extract_content_from_uploadfiles(data_files, "data", image_counter)
    theory_text, theory_images, theory_assets, theory_tables, image_counter = await extract_content_from_uploadfiles(theory_files, "theory", image_counter)
    waveforms_text, waveforms_images, waveforms_assets, waveforms_tables, image_counter = await extract_content_from_uploadfiles(waveforms_files, "waveforms", image_counter)
    procedure_text, procedure_images, procedure_assets, procedure_tables, image_counter = await extract_content_from_uploadfiles(procedure_files, "procedure", image_counter)
    conclusion_text, conclusion_images, conclusion_assets, conclusion_tables, image_counter = await extract_content_from_uploadfiles(conclusion_files, "conclusion", image_counter)
    _, schema_images_list, schema_assets, schema_tables, image_counter = await extract_content_from_uploadfiles(schema_files, "schema", image_counter)

    if not schema_images_list and fallback_schema_images:
        schema_images_list = fallback_schema_images
        schema_assets.extend(fallback_schema_assets)

    if not waveforms_images and fallback_waveforms_images:
        waveforms_images = fallback_waveforms_images
        waveforms_assets.extend(fallback_waveforms_assets)

    (
        data_images,
        data_assets,
        waveforms_images,
        waveforms_assets,
        data_text,
        waveforms_text,
    ) = _reroute_xlsx_charts_from_data_to_waveforms(
        data_images=data_images,
        data_assets=data_assets,
        waveforms_images=waveforms_images,
        waveforms_assets=waveforms_assets,
        data_text=data_text,
        waveforms_text=waveforms_text,
    )

    if not procedure_text.strip():
        procedure_section = assignment_sections.get("procedure") if assignment_sections else None
        if isinstance(procedure_section, dict):
            procedure_text = str(procedure_section.get("text", "")).strip()

    assignment_theory_text = ""
    assignment_conclusion_text = ""
    if assignment_sections:
        theory_section = assignment_sections.get("theory")
        if isinstance(theory_section, dict):
            assignment_theory_text = str(theory_section.get("text", "")).strip()

        conclusion_section = assignment_sections.get("conclusion")
        if isinstance(conclusion_section, dict):
            assignment_conclusion_text = str(conclusion_section.get("text", "")).strip()

    all_assets = (
        assignment_assets + instruments_assets + data_assets + theory_assets +
        waveforms_assets + procedure_assets + conclusion_assets + schema_assets
    )
    image_registry = {asset.image_id: asset for asset in all_assets}
    image_catalog_text = "\n".join(
        f"- {asset.image_id}: {asset.filename} (sekce: {asset.section})"
        for asset in all_assets
    )

    inputs_map = {
        "assignment_text": assignment_text,
        "instruments_text": instruments_text,
        "data_text": data_text,
    "data_tables": data_tables,
        "theory_text": theory_text,
    "assignment_theory_text": assignment_theory_text,
        "waveforms_text": waveforms_text,
        "procedure_text": procedure_text,
        "conclusion_text": conclusion_text,
    "assignment_conclusion_text": assignment_conclusion_text,
        "schema_images": schema_images_list,
        "schema_image_ids": [asset.image_id for asset in schema_assets],
        "data_images": data_images,
        "data_image_ids": [asset.image_id for asset in data_assets],
        "waveforms_images": waveforms_images,
        "waveforms_image_ids": [asset.image_id for asset in waveforms_assets],
    "report_scope": resolved_report_scope,
        "username": username,
        "topic": topic,
        "image_catalog_text": image_catalog_text,
        "images_lists": [
            assignment_images,
            instruments_images,
            data_images,
            theory_images,
            waveforms_images,
            procedure_images,
            conclusion_images,
        ],
        "section_tables": {
            "assignment": assignment_tables,
            "instruments": instruments_tables,
            "data": data_tables,
            "theory": theory_tables,
            "waveforms": waveforms_tables,
            "procedure": procedure_tables,
            "conclusion": conclusion_tables,
            "schema": schema_tables,
        },
    }

    try:
        ai_data = generate_lab_report_structured(
            api_key=resolved_api_key,
            model_name=model_name,
            topic=topic,
            inputs_map=inputs_map,
            is_handwritten=is_handwritten,
            report_scope=resolved_report_scope,
        )
    except Exception as e:
        return JSONResponse(status_code=500, content={"error": f"Generování selhalo: {str(e)}"})

    quality = run_quality_check(ai_data, image_registry)
    if quality.status == "FAIL":
        return JSONResponse(
            status_code=422,
            content={
                "error": "Quality gate FAIL",
                "issues": [issue.model_dump() for issue in quality.issues],
            },
        )

    template_path = os.path.join(BASE_DIR, "DOCS", "sablona_elab_em.docx")
    docx_buffer = fill_template_docx_structured(
        template_path=template_path,
        inputs_map=inputs_map,
        ai_content=ai_data,
        image_registry=image_registry,
        username=username,
        topic=topic,
    )
    docx_buffer.seek(0)

    headers = {
        "Content-Disposition": 'attachment; filename="laboratorni_protokol.docx"',
        "X-Quality-Status": quality.status,
    }
    return StreamingResponse(
        docx_buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers,
    )


if __name__ == "__main__":
    host = os.getenv("FASTAPI_HOST", "127.0.0.1")
    port = int(os.getenv("FASTAPI_PORT", "8000"))
    uvicorn.run("fastapi_server:app", host=host, port=port, reload=False)
